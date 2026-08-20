"""docx <-> pdf conversion.

docx -> pdf : docx2pdf
pdf  -> docx: pdfplumber

Output no longer lives on disk after conversion: finalize_output() stores
the result as a DB blob and returns a download token/url, which is attached
to the ConversionResult so the agent can hand it back to the user.
"""
import sys
import logging
import warnings
from tools.base import ConversionResult
from utils.errors import EngineError
from utils.file_io import validate_input_path, resolve_output_path, finalize_output

logging.basicConfig(stream=sys.stderr)
warnings.filterwarnings("ignore", module="fitz")
warnings.filterwarnings("ignore", category=DeprecationWarning)


def _attach_download(result: ConversionResult, dest) -> ConversionResult:
    # ConversionResult.to_json() only serializes success/output_path/
    # input_format/output_format/size_bytes/meta/error — plain attributes
    # set after construction (e.g. result.download_url = ...) are silently
    # dropped. `meta` is the real extension point; it's serialized.
    info = finalize_output(dest)
    result.meta["download_token"] = info["token"]
    result.meta["download_url"] = info["url"]
    result.meta["download_filename"] = info["filename"]

    # output_path currently points at `dest`, which finalize_output() just
    # deleted — that path is dead the moment this function returns. Left
    # as-is, the LLM (which naturally narrates the most prominent field)
    # tends to report that stale local disk path to the user instead of
    # the working download link buried in meta. Overwrite it with the
    # actual reachable location so whatever the agent quotes back is
    # correct by construction, not by hoping it picks the right field.
    result.output_path = info["url"]
    return result


def docx_to_pdf(input_path: str, output_path: str | None = None) -> ConversionResult:
    src = validate_input_path(input_path, {"docx", "doc"})
    dest = resolve_output_path(output_path, src, "pdf")
    dest.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx2pdf import convert
    except ImportError as e:
        raise EngineError("docx2pdf not installed. Run: pip install docx2pdf") from e

    try:
        convert(str(src), str(dest))
    except Exception as e:  # noqa: BLE001
        raise EngineError(f"docx2pdf conversion failed: {e}") from e

    if not dest.exists():
        raise EngineError("docx2pdf reported success but no output file was produced")

    result = ConversionResult.ok(dest, input_format=src.suffix.lstrip("."), output_format="pdf")
    return _attach_download(result, dest)


def _pdf_to_docx_via_pdfplumber(input_path: str, output_path: str | None = None) -> ConversionResult:
    src = validate_input_path(input_path, {"pdf"})
    dest = resolve_output_path(output_path, src, "docx")
    dest.parent.mkdir(parents=True, exist_ok=True)

    try:
        import pdfplumber
        from docx import Document
    except ImportError as e:
        raise EngineError("pdfplumber/python-docx not installed. Run: pip install pdfplumber python-docx") from e

    try:
        doc = Document()
        with pdfplumber.open(str(src)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

                for table in page.extract_tables():
                    if not table:
                        continue
                    rows, cols = len(table), len(table[0])
                    docx_table = doc.add_table(rows=rows, cols=cols)
                    for r, row in enumerate(table):
                        for c, cell in enumerate(row):
                            docx_table.cell(r, c).text = cell or ""

                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()

        doc.save(str(dest))
    except Exception as e:  # noqa: BLE001
        raise EngineError(f"pdfplumber conversion failed: {e}") from e

    if not dest.exists():
        raise EngineError("pdfplumber reported success but no output file was produced")

    result = ConversionResult.ok(dest, input_format="pdf", output_format="docx")
    return _attach_download(result, dest)